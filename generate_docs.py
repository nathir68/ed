import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = docx.Document()

    # Title
    title = doc.add_heading('EduConnect: Comprehensive Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Ideology
    doc.add_heading('1. Ideology', level=1)
    doc.add_paragraph(
        "EduConnect was conceived with a clear ideology: to bridge the gap between traditional learning environments and modern digital accessibility. Our goal is to empower educators with seamless tools to distribute knowledge and provide students with a unified, distraction-free platform for accessing study materials, attending virtual classes, and collaborating with their peers. By prioritizing a clean user interface, intuitive workflows, and robust functionality, EduConnect aims to democratize education technology for institutions of all sizes."
    )

    # 2. What we have used in this (Tech Stack)
    doc.add_heading('2. What we have used in this (Tech Stack)', level=1)
    tech_stack = doc.add_paragraph()
    tech_stack.add_run("Frontend: ").bold = True
    tech_stack.add_run("HTML5, CSS3, Vanilla JavaScript (Jinja2 Templates)\n")
    tech_stack.add_run("Backend: ").bold = True
    tech_stack.add_run("Python 3, Flask (Web Framework)\n")
    tech_stack.add_run("Database: ").bold = True
    tech_stack.add_run("MySQL (Relational Database Management System)\n")
    tech_stack.add_run("Real-Time Communication: ").bold = True
    tech_stack.add_run("Flask-SocketIO, WebRTC (for live video sessions and chat)\n")
    tech_stack.add_run("Authentication & Security: ").bold = True
    tech_stack.add_run("Werkzeug Security (Password Hashing), UUIDs, OTP Generation")

    # 3. Features
    doc.add_heading('3. Features', level=1)
    features = [
        "Multi-Role Authentication: Secure login and signup for Super Admins, College Admins, Teachers, and Students.",
        "Role-Based Dashboards: Customized views and privileges based on the user's role.",
        "Resource Management: Teachers can upload notes, past papers, and study materials (PDF, DOCX, images) directly to their department and year.",
        "Live Virtual Classes: Integrated video conferencing using WebRTC, complete with screen sharing, camera toggles, and real-time chat.",
        "Search & Filtering: Advanced search capabilities allowing students to instantly filter resources and live sessions by name or topic.",
        "In-App Notifications: Real-time alerts (with a bell icon) notifying students when new materials are uploaded or classes are scheduled.",
        "Attendance Tracking: Automated and manual attendance logs for teachers to monitor student participation.",
        "Dark Mode Toggle: A built-in theme switcher for optimal viewing comfort in low-light environments.",
        "Advanced Security: OTP-based forgot password flows, email verification prompts, and secure password hashing."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # 4. How it works
    doc.add_heading('4. How it works', level=1)
    doc.add_paragraph(
        "EduConnect operates on a hierarchical college structure. A Super Admin creates institutions (Colleges) and appoints College Admins. College Admins then structure their institution by adding Departments and assigning Teachers and Students to specific years of study."
    )
    doc.add_paragraph(
        "When a Teacher logs in, they are presented with a dashboard where they can upload resources or schedule live sessions targeting specific departments and years. These actions trigger backend database updates and automatically generate notifications."
    )
    doc.add_paragraph(
        "When a Student logs in, the system dynamically filters the database to show only the resources and upcoming classes relevant to their specific department and year of study. They can join live sessions via the built-in WebRTC video interface, where attendance is logged automatically."
    )

    # 5. System Architecture
    doc.add_heading('5. System Architecture', level=1)
    doc.add_paragraph(
        "The application follows a standard Model-View-Controller (MVC) architectural pattern, adapted for Flask:"
    )
    arch = doc.add_paragraph()
    arch.add_run("Model (Database): ").bold = True
    arch.add_run("A MySQL relational database handles data persistence. Core tables include users, colleges, departments, resources, classes, attendance, notifications, and password_resets. Foreign keys ensure referential integrity.\n")
    
    arch.add_run("View (Frontend): ").bold = True
    arch.add_run("Jinja2 templates dynamically render HTML pages. A global layout (base.html) ensures consistent navigation, footers, and flash messaging across the platform.\n")
    
    arch.add_run("Controller (Backend Routes): ").bold = True
    arch.add_run("The Flask application (app.py) processes incoming HTTP requests, interacts with the MySQL database via Python's mysql-connector, applies business logic (like verifying OTPs or checking user roles), and returns the appropriate View.\n")
    
    arch.add_run("Real-Time Engine: ").bold = True
    arch.add_run("A parallel WebSocket server (Flask-SocketIO) runs alongside the HTTP server to handle low-latency events like chat messages and WebRTC signaling for the video classrooms.")

    # Save
    doc.save('EduConnect_Documentation.docx')
    print("Documentation generated successfully.")

if __name__ == '__main__':
    create_doc()
